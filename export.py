import pandas as pd
import streamlit as st
from io import BytesIO, StringIO
from docx import Document
from docx.shared import RGBColor
from xhtml2pdf import pisa
from html import escape
import difflib

def export_csv(results):
    df = pd.DataFrame(results)
    csv = df[["class", "name", "type", "summary"]].to_csv(index=False).encode("utf-8")
    st.download_button("🗕️ Download CSV", data=csv, file_name="servicenow_records.csv", mime="text/csv")

def export_pdf(results, context=5):
    html = """
    <html><head>
    <style>
    body { font-family: Arial; }
    h3 { color: #333; }
    pre { background: #f4f4f4; padding: 10px; }
    .added { background-color: #d4fcdc; }
    .removed { background-color: #fddede; }
    .context { color: #888; }
    </style>
    </head><body><h2>ServiceNow Record Summary</h2>
    """
    for r in results:
        html += f"<h3>{r['name']} ({r['type']})</h3>"
        html += f"<p><b>AI Summary:</b> {r['summary']}</p><pre>"

        if 'before' in r and r['before']:
            before_lines = r['before'].splitlines()
            after_lines = r['after'].splitlines()
            matcher = difflib.SequenceMatcher(None, before_lines, after_lines)

            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'equal':
                    continue

                start_before = max(i1 - context, 0)
                end_after = min(j2 + context, len(after_lines))

                html += "<span class='context'>Context Before:\n"
                html += "\n".join(f"{idx+1:4}: {escape(before_lines[idx])}" for idx in range(start_before, i1))
                html += "\n</span>"

                html += "<span class='removed'>"
                html += "\n".join(f"{idx+1:4}: - {escape(before_lines[idx])}" for idx in range(i1, i2))
                html += "\n</span>"

                html += "<span class='added'>"
                html += "\n".join(f"{idx+1:4}: + {escape(after_lines[idx])}" for idx in range(j1, j2))
                html += "\n</span>"

                html += "<span class='context'>Context After:\n"
                html += "\n".join(f"{idx+1:4}: {escape(after_lines[idx])}" for idx in range(j2, end_after))
                html += "\n</span>"

                html += "-"*40 + "\n"
        else:
            html += escape(r['after'][:2000])

        html += "</pre><hr>"
    html += "</body></html>"
    output = BytesIO()
    pisa.CreatePDF(StringIO(html), dest=output)
    st.download_button("📄 Download PDF", data=output.getvalue(), file_name="servicenow_summary.pdf", mime="application/pdf")

def export_docx(results, context=5):
    doc = Document()
    doc.add_heading('🧠 ServiceNow Record Summary', 0)

    for r in results:
        doc.add_heading(f"{r['name']} ({r['type']})", level=1)
        doc.add_paragraph(f"AI Summary: {r['summary']}")

        if 'before' in r and r['before']:
            before_lines = r['before'].splitlines()
            after_lines = r['after'].splitlines()
            matcher = difflib.SequenceMatcher(None, before_lines, after_lines)

            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'equal':
                    continue

                start_before = max(i1 - context, 0)
                end_after = min(j2 + context, len(after_lines))

                doc.add_paragraph("Context Before:", style='Intense Quote')
                for idx in range(start_before, i1):
                    doc.add_paragraph(f"{idx+1:4}: {before_lines[idx]}")

                doc.add_paragraph("Changes:", style='Intense Quote')
                for idx in range(i1, i2):
                    p = doc.add_paragraph()
                    run = p.add_run(f"{idx+1:4}: - {before_lines[idx]}")
                    run.font.color.rgb = RGBColor(255, 0, 0)

                for idx in range(j1, j2):
                    p = doc.add_paragraph()
                    run = p.add_run(f"{idx+1:4}: + {after_lines[idx]}")
                    run.font.color.rgb = RGBColor(0, 128, 0)

                doc.add_paragraph("Context After:", style='Intense Quote')
                for idx in range(j2, end_after):
                    doc.add_paragraph(f"{idx+1:4}: {after_lines[idx]}")

                doc.add_paragraph("\n" + "-"*40 + "\n")
        else:
            doc.add_paragraph(r['after'][:2000])

    output = BytesIO()
    doc.save(output)
    st.download_button("📄 Download DOCX", data=output.getvalue(), file_name="servicenow_summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def export_html(results, context=5):
    html = """<html><head><style>
    .added { background-color: #d4fcdc; }
    .removed { background-color: #fddede; }
    .context { color: #888; }
    pre { background: #f4f4f4; padding: 10px; border-radius: 5px; }
    </style></head><body><h2>🧠 ServiceNow Record Summary</h2>"""

    for r in results:
        html += f"<h3>{r['name']} ({r['type']})</h3>"
        html += f"<p><strong>AI Summary:</strong> {r['summary']}</p><pre>"

        if 'before' in r and r['before']:
            before_lines = r['before'].splitlines()
            after_lines = r['after'].splitlines()
            matcher = difflib.SequenceMatcher(None, before_lines, after_lines)

            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'equal':
                    continue

                start_before = max(i1 - context, 0)
                end_after = min(j2 + context, len(after_lines))

                html += "<span class='context'>Context Before:\n"
                html += "\n".join(f"{idx+1:4}: {escape(before_lines[idx])}" for idx in range(start_before, i1))
                html += "\n</span>"

                html += "<span class='removed'>"
                html += "\n".join(f"{idx+1:4}: - {escape(before_lines[idx])}" for idx in range(i1, i2))
                html += "\n</span>"

                html += "<span class='added'>"
                html += "\n".join(f"{idx+1:4}: + {escape(after_lines[idx])}" for idx in range(j1, j2))
                html += "\n</span>"

                html += "<span class='context'>Context After:\n"
                html += "\n".join(f"{idx+1:4}: {escape(after_lines[idx])}" for idx in range(j2, end_after))
                html += "\n</span>"

                html += "-"*40 + "\n"
        else:
            html += escape(r['after'][:2000])

        html += "</pre><hr>"
    html += "</body></html>"

    st.download_button("🌐 Download HTML", data=html, file_name="servicenow_summary.html", mime="text/html")
